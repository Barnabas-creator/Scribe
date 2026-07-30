"""DocModel -> .docx using python-docx for layout, pandoc/OMML for formulas.

python-docx owns paragraph/heading/list/table/image structure; each formula is
turned into a native Word equation via omml.latex_to_omml. Anything that fails
or is low-confidence is highlighted yellow and collected into a review list so a
human can spot-check it against the original scan.

MinerU does not emit per-element confidence, so by default every non-trivial
formula is added to the review list (flag_formulas=True): formulas are the most
error-prone part of exam OCR and warrant a human glance.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from lxml import etree

from . import model as M
from .images import add_image
from .omml import FormulaConversionError, batch_latex_to_omml, latex_to_omml

LOW_CONFIDENCE = 0.6
# Single-symbol inline formulas are near-certain and would drown the report.
_TRIVIAL_FORMULA = re.compile(r"^[A-Za-z\\]{1,7}(_\{?\w\}?)?$")


def _worth_review(latex: str) -> bool:
    return not _TRIVIAL_FORMULA.match(latex.strip())


@dataclass
class ReviewItem:
    page: int
    bbox: tuple
    kind: str  # "formula" | "text" | "image"
    detail: str
    reason: str  # "formula_check" | "low_confidence" | "conversion_failed"
    crop: str = ""  # base64 PNG of the source region, for side-by-side review


@dataclass
class _Ctx:
    review: list
    flag_formulas: bool
    # Document-wide batch conversion result: (latex, block) -> serialized bytes.
    # None means the batch missed that slot and per-formula conversion applies.
    precomputed: dict = field(default_factory=dict)
    # Inline formulas in one paragraph share a single crop; only the first
    # review item carries the image to keep the report small.
    cropped: set = field(default_factory=set)

    def crop_once(self, blk: M.Block) -> str:
        if not blk.crop_b64 or id(blk) in self.cropped:
            return ""
        self.cropped.add(id(blk))
        return blk.crop_b64


def _set_default_fonts(doc, cjk: str = "宋体", latin: str = "Times New Roman") -> None:
    """Set document default fonts: Latin -> Times New Roman, CJK -> 宋体.

    Word picks the font per character run: Western chars use w:ascii/w:hAnsi,
    Chinese chars use w:eastAsia. Applied to Normal + Title + Heading styles so
    body text and headings all follow the convention. Formulas stay in Cambria
    Math (untouched, as math should).
    """
    targets = ["Normal", "Title"] + [f"Heading {i}" for i in range(1, 10)]
    for name in targets:
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.name = latin
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
        rfonts.set(qn("w:cs"), latin)
        rfonts.set(qn("w:eastAsia"), cjk)


def _collect_formulas(doc_model: M.DocModel) -> list[tuple[str, bool]]:
    """Collect every formula in render order for one batch conversion."""
    items: list[tuple[str, bool]] = []
    for blk in doc_model.sorted_blocks():
        if blk.type == M.FORMULA:
            items.append((blk.latex.strip(), blk.block_formula))
        elif blk.type in M.TEXT_TYPES:
            items.extend((s.latex.strip(), False) for s in blk.spans if s.is_formula)
        elif blk.type == M.TABLE:
            items.extend((s.latex.strip(), False)
                         for row in blk.rows for cell in row for s in cell if s.is_formula)
    return items


# Control characters XML 1.0 rejects (tab and newline kept). PDF text layers do
# contain these -- private-use code points and hyphenation marks extract as
# \x02 or \x0c -- and python-docx refuses to write them.
_XML_BAD = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def _scrub(doc_model: M.DocModel) -> None:
    """Strip illegal control characters from every string bound for the docx.

    Done once at the render entry point so headings, body, tables and formulas
    are all covered; missing one call site would fail the whole document.
    """
    for blk in doc_model.blocks:
        if blk.latex:
            blk.latex = _XML_BAD.sub("", blk.latex)
        for span in blk.spans:
            if span.text:
                span.text = _XML_BAD.sub("", span.text)
        for row in blk.rows or []:
            for cell in row:
                for span in cell:
                    if span.text:
                        span.text = _XML_BAD.sub("", span.text)


def render(doc_model: M.DocModel, out_path: str, flag_formulas: bool = True) -> list[ReviewItem]:
    doc = Document()
    _set_default_fonts(doc)
    _scrub(doc_model)
    # One pandoc run for the whole document; unmatched slots fall back to
    # per-formula conversion, which itself degrades to highlighted text.
    precomputed: dict = {}
    formulas = _collect_formulas(doc_model)
    for key, el in zip(formulas, batch_latex_to_omml(formulas)):
        precomputed.setdefault(key, None if el is None else etree.tostring(el))
    ctx = _Ctx(review=[], flag_formulas=flag_formulas, precomputed=precomputed)
    for blk in doc_model.sorted_blocks():
        _render_block(doc, blk, ctx)
    doc.save(out_path)
    return ctx.review


def _render_block(doc, blk: M.Block, ctx: _Ctx) -> None:
    if blk.type == M.HEADING:
        p = doc.add_heading("", level=max(1, min(blk.level, 9)))
        _render_spans(p, blk, ctx)
    elif blk.type == M.PARAGRAPH:
        _render_spans(doc.add_paragraph(), blk, ctx)
    elif blk.type == M.LIST_ITEM:
        _render_spans(doc.add_paragraph(style="List Bullet"), blk, ctx)
    elif blk.type == M.FORMULA:
        _add_formula(doc.add_paragraph(), blk.latex, blk.block_formula, blk, ctx)
    elif blk.type == M.PICTURE:
        p = doc.add_paragraph()
        if not add_image(p, blk.image_path):
            ctx.review.append(ReviewItem(blk.page, blk.bbox, "image", blk.image_path, "conversion_failed"))
    elif blk.type == M.TABLE:
        _render_table(doc, blk, ctx)


def _render_spans(paragraph, blk: M.Block, ctx: _Ctx) -> None:
    low = blk.confidence < LOW_CONFIDENCE
    for span in blk.spans:
        if span.is_formula:
            _add_formula(paragraph, span.latex, False, blk, ctx)
        else:
            run = paragraph.add_run(span.text)
            run.bold = span.bold
            run.italic = span.italic
            if low:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    if low and any(not s.is_formula for s in blk.spans):
        ctx.review.append(ReviewItem(
            blk.page, blk.bbox, "text",
            "".join(s.text for s in blk.spans)[:60], "low_confidence"))


def _add_formula(paragraph, latex: str, block: bool, blk: M.Block, ctx: _Ctx) -> None:
    try:
        cached = ctx.precomputed.get((latex.strip(), block))
        if cached is not None:
            paragraph._p.append(etree.fromstring(cached))
        else:  # batch missed this slot: convert individually
            paragraph._p.append(latex_to_omml(latex, block=block))
        if blk.confidence < LOW_CONFIDENCE:
            ctx.review.append(ReviewItem(blk.page, blk.bbox, "formula", latex, "low_confidence", crop=ctx.crop_once(blk)))
        elif ctx.flag_formulas and _worth_review(latex):
            ctx.review.append(ReviewItem(blk.page, blk.bbox, "formula", latex, "formula_check", crop=ctx.crop_once(blk)))
    except FormulaConversionError:
        run = paragraph.add_run(f" [公式待校对: {latex}] ")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        ctx.review.append(ReviewItem(blk.page, blk.bbox, "formula", latex, "conversion_failed", crop=ctx.crop_once(blk)))


def _table_summary(blk: M.Block) -> str:
    """Text summary of a table for the review report."""
    head = " | ".join("".join(s.text for s in cell).strip() or "·" for cell in blk.rows[0])
    body = "\n".join(
        " | ".join("".join(s.text for s in cell).strip() or "·" for cell in row)
        for row in blk.rows[1:]
    )
    return f"{len(blk.rows)} 行 × {max(len(r) for r in blk.rows)} 列\n{head}\n{body}".strip()


def _render_table(doc, blk: M.Block, ctx: _Ctx) -> None:
    if not blk.rows:
        return
    # Tables need review too: shifted rows and misread digits hide better than
    # formula errors.
    if ctx.flag_formulas:
        ctx.review.append(ReviewItem(blk.page, blk.bbox, "table", _table_summary(blk),
                                     "table_check", crop=ctx.crop_once(blk)))
    n_cols = max(len(r) for r in blk.rows)
    table = doc.add_table(rows=len(blk.rows), cols=n_cols)
    table.style = "Table Grid"
    for i, row in enumerate(blk.rows):
        for j, cell_spans in enumerate(row):
            cell = table.cell(i, j)
            p = cell.paragraphs[0]
            p.text = ""
            for span in cell_spans:
                if span.is_formula:
                    _add_formula(p, span.latex, False, blk, ctx)
                else:
                    run = p.add_run(span.text)
                    run.bold = span.bold
                    run.italic = span.italic
