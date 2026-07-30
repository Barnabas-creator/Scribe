"""Image handling: embed extracted figures into docx, and crop formula
snapshots from the source file for the review report.

MinerU extracts figures as standalone files, but equation blocks carry only a
bbox and LaTeX -- no img_path -- so formula crops are cut here with PyMuPDF.
"""

from __future__ import annotations

import base64
from pathlib import Path

from docx.shared import Emu

from . import model as M

try:
    from PIL import Image
except Exception:  # pragma: no cover - Pillow is a hard dep, guarded for safety
    Image = None

EMU_PER_INCH = 914400
EMU_PER_PX_96DPI = 9525  # 1 px at 96 dpi
DEFAULT_MAX_WIDTH_EMU = int(6.0 * EMU_PER_INCH)  # usable width on Letter/A4 w/ margins


def add_image(paragraph, path: str, max_width_emu: int = DEFAULT_MAX_WIDTH_EMU) -> bool:
    """Add image at `path` inline into `paragraph`. Returns True on success."""
    run = paragraph.add_run()
    width_emu = None
    if Image is not None:
        try:
            with Image.open(path) as im:
                w_px = im.size[0]
            width_emu = min(w_px * EMU_PER_PX_96DPI, max_width_emu)
        except Exception:
            width_emu = None
    try:
        if width_emu:
            run.add_picture(path, width=Emu(width_emu))
        else:
            run.add_picture(path)
        return True
    except Exception:
        run.add_text(f"[图片缺失或无法读取: {path}]")
        return False


# MinerU normalizes bboxes to 0-1000, y-down (same axis as fitz, no flip needed).
_BBOX_SCALE = 1000.0


def attach_formula_crops(source: str | Path, doc_model: M.DocModel,
                         zoom: float = 2.0, pad: float = 2.0) -> int:
    """Attach a base64 PNG crop of the original region to formula-bearing
    blocks; returns how many succeeded.

    Display formulas are cropped exactly; inline formulas have no bbox of their
    own (MinerU reports paragraph granularity), so the whole paragraph is cut.
    Failures leave the field empty -- the report tolerates missing crops.
    """
    blocks = [b for b in doc_model.blocks
              if any(b.bbox) and (b.type == M.FORMULA
                                  or (b.type in M.TEXT_TYPES
                                      and any(s.is_formula for s in b.spans)))]
    if not blocks:
        return 0

    import pymupdf

    n = 0
    try:
        doc = pymupdf.open(str(source))
    except Exception:
        return 0
    with doc:
        for b in blocks:
            try:
                page = doc[max(0, min(b.page - 1, doc.page_count - 1))]
                w, h = page.rect.width, page.rect.height
                x0, y0, x1, y1 = b.bbox
                rect = pymupdf.Rect(x0 / _BBOX_SCALE * w - pad, y0 / _BBOX_SCALE * h - pad,
                                    x1 / _BBOX_SCALE * w + pad, y1 / _BBOX_SCALE * h + pad)
                pix = page.get_pixmap(matrix=pymupdf.Matrix(zoom, zoom), clip=rect, alpha=False)
                b.crop_b64 = base64.b64encode(pix.tobytes("png")).decode()
                n += 1
            except Exception:
                continue
    return n
