"""Smoke test: prove pandoc produces editable Word formulas (no macro needed).

Builds a docx mixing Chinese text, inline formula, display formula, a bold
run, and a deliberately broken formula (to exercise the fallback). Then it
re-opens the file and asserts native <m:oMath> elements are present.
"""

import sys
from pathlib import Path
import zipfile

sys.path.insert(0, str(__import__("pathlib").Path(__file__).resolve().parents[1] / "src"))

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from p2w.omml import latex_to_omml, FormulaConversionError

TMP = Path(__file__).resolve().parents[1] / "tmp"
TMP.mkdir(exist_ok=True)
OUT = str(TMP / "verify_omml.docx")


def add_formula(paragraph, latex, block=False):
    try:
        paragraph._p.append(latex_to_omml(latex, block=block))
        return True
    except FormulaConversionError:
        run = paragraph.add_run(f" [公式待修: {latex}] ")
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        return False


doc = Document()
doc.add_heading("一、选择题", level=1)

p = doc.add_paragraph("1. 已知导数定义 ")
add_formula(p, r"\frac{f(x+h)-f(x)}{h}", block=False)
p.add_run(" ，当 ")
add_formula(p, r"h \to 0", block=False)
p.add_run(" 时求极限。其中 ")
p.add_run("重点").bold = True
p.add_run(" 是收敛性。")

doc.add_paragraph("下面是独立成行的二次方程求根公式：")
pb = doc.add_paragraph()
add_formula(pb, r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}", block=True)

doc.add_paragraph("矩阵与分段函数：")
pm = doc.add_paragraph()
add_formula(pm, r"\begin{cases} x^2 & x>0 \\ -x & x\le 0 \end{cases}", block=True)

p3 = doc.add_paragraph("故意写坏的公式（测试降级）：")
add_formula(p3, r"\frac{1}{", block=False)

doc.save(OUT)

# Re-open and count native math elements
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
with zipfile.ZipFile(OUT) as zf:
    xml = zf.read("word/document.xml").decode("utf-8")
count = xml.count("<m:oMath")
print(f"saved: {OUT}")
print(f"native <m:oMath...> elements in document.xml: {count}")
print("contains highlighted fallback text:", "公式待修" in xml)
assert count >= 4, "expected at least 4 native math elements"
print("PASS: formulas are native editable Word equations, not images/text")

# ---- Batch conversion: one pandoc run for the whole document ----
from unittest import mock

from lxml import etree as _etree

from p2w.omml import batch_latex_to_omml

mixed = [
    (r"\frac{f(x+h)-f(x)}{h}", False),
    (r"x = \frac{-b \pm \sqrt{b^2 - 4ac}}{2a}", True),
    (r"h \to 0", False),
    (r"\frac{1}{", False),  # 故意写坏：这一槽必须是 None，走逐条兜底
    (r"\begin{cases} x^2 & x>0 \\ -x & x\le 0 \end{cases}", True),
]
batched = batch_latex_to_omml(mixed)
assert len(batched) == len(mixed), "返回列表必须和输入一一对齐"
assert batched[3] is None, "坏公式不该出 OMML"
for i, el in enumerate(batched):
    if i == 3:
        continue
    assert el is not None, f"slot {i} 丢了"
    want = "oMathPara" if mixed[i][1] else "oMath"
    assert _etree.QName(el).localname == want, f"slot {i} 结构不对"
    # Batch output must match per-formula output byte for byte
    assert _etree.tostring(el) == _etree.tostring(latex_to_omml(*mixed[i])), f"slot {i} 与逐条结果不一致"
print("PASS: 批量转换对齐正确，且与逐条结果逐字节一致")

# Empty input must not invoke pandoc at all
with mock.patch("p2w.omml.subprocess.run", side_effect=AssertionError("不该调 pandoc")):
    assert batch_latex_to_omml([]) == []
print("PASS: 空列表跳过 pandoc")

# Batch failure yields all-None slots so the caller can fall back
with mock.patch("p2w.omml.subprocess.run", side_effect=FileNotFoundError("pandoc")):
    fallen = batch_latex_to_omml(mixed)
assert fallen == [None] * len(mixed), fallen
print("PASS: 批量失败全槽位 None，可回退逐条转换")
